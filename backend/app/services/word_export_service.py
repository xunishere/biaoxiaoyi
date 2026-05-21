"""Word 导出服务 — 纯标题层级 + 目录 + 可配字体格式。"""

import io
import re
from urllib.parse import quote

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING

# ── 段落格式 ──
def _format_para(para: docx.text.paragraph.Paragraph) -> None:
    """统一段落格式：首行缩进2字符。"""
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)

from ..models.schemas import WordExportOutlineItem, WordExportRequest

# ── 可配置字体格式 ──
FONT_NAME = "宋体"                  # 中文字体
FONT_SIZES = {
    "heading1": Pt(16),             # 一级标题
    "heading2": Pt(14),             # 二级标题
    "heading3": Pt(13),             # 三级标题
    "heading4": Pt(12),             # 四级标题（加粗段落）
    "body":     Pt(12),             # 正文
}
TOC_TITLE = "目  录"                # 目录标题
TOC_LEVELS = "1-4"                  # 目录包含的标题级别
# ──────────────────────


def _set_run_font(run: docx.text.run.Run, size: Pt | None = None) -> None:
    """设置 run 的字体和字号。"""
    run.font.name = FONT_NAME
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        rpr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size


def _add_toc(doc: docx.Document) -> None:
    """在文档开头插入 Word 目录域，Word 打开后自动生成带页码的目录。"""
    # 目录标题
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run(TOC_TITLE)
    run.bold = True
    _set_run_font(run, Pt(14))

    # TOC 域代码
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "{TOC_LEVELS}" \\h </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar_end)

    # 分页
    doc.add_page_break()


class WordExportService:
    """负责将目录数据导出为 Word 文档。"""

    @staticmethod
    def export_outline(request: WordExportRequest) -> tuple[io.BytesIO, dict[str, str]]:
        doc = docx.Document()
        WordExportService._init_styles(doc)
        _add_toc(doc)
        WordExportService._add_outline_items(doc, request.outline)

        # 标记：Word 打开时自动更新域（目录自动生成）
        settings = doc.settings.element
        update_fields = parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
        settings.append(update_fields)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{request.project_name or '标书文档'}.docx"
        headers = {
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}",
        }
        return buffer, headers

    @staticmethod
    def _init_styles(doc: docx.Document) -> None:
        """初始化标题和正文样式。"""
        try:
            style_map = {
                "Heading 1": FONT_SIZES["heading1"],
                "Heading 2": FONT_SIZES["heading2"],
                "Heading 3": FONT_SIZES["heading3"],
                "Heading 4": FONT_SIZES["heading4"],
                "Normal":    FONT_SIZES["body"],
            }
            for name, size in style_map.items():
                if name not in doc.styles:
                    # 创建缺失的样式（如 Heading 4）
                    if name.startswith("Heading"):
                        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                if name not in doc.styles:
                    continue
                style = doc.styles[name]
                style.font.name = FONT_NAME
                style.font.size = size
                if style._element.rPr is None:
                    style._element._add_rPr()
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            # Normal 不加粗
            if "Normal" in doc.styles:
                doc.styles["Normal"].font.bold = False
        except Exception:
            pass

    @staticmethod
    def _add_markdown_runs(para: docx.text.paragraph.Paragraph, text: str) -> None:
        """行内格式：**加粗** *斜体* `代码`。"""
        pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
        parts = re.split(pattern, text)
        for part in parts:
            if not part:
                continue
            run = para.add_run()
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                run.text = part[1:-1]
            else:
                run.text = part
            _set_run_font(run, FONT_SIZES["body"])

    @staticmethod
    def _add_body_paragraph(doc: docx.Document, text: str) -> None:
        """添加正文段落（清理markdown、空格、破折号，统一格式）。"""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = text.replace('*', '')
        text = text.strip()  # 去前后空格
        text = re.sub(r'^[\s　]+', '', text)  # 去前导空格/全角空格
        text = text.replace('——', '-')  # 破折号→短横
        para = doc.add_paragraph()
        WordExportService._add_markdown_runs(para, text)
        _format_para(para)

    @staticmethod
    def _add_body_content(doc: docx.Document, content: str) -> None:
        """输出正文：识别表格和段落。"""
        blocks = content.split("\n\n")
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            lines = [ln for ln in block.split("\n") if ln.strip()]
            # 表格内容跳过
            if len(lines) >= 2 and all("|" in ln for ln in lines):
                continue
            for line in lines:
                line = re.sub(r'^#{1,6}\s+', '', line)
                line = re.sub(r'^[-*+]\s+', '', line)
                line = re.sub(r'^\d+[\.\)]\s*', '', line)
                WordExportService._add_body_paragraph(doc, line)

    @staticmethod
    def _add_table_from_markdown(doc: docx.Document, lines: list[str]) -> None:
        """将markdown表格行渲染为Word表格。异常行退回段落。"""
        rows = []
        overflow_lines: list[str] = []
        for ln in lines:
            cells = [c.strip() for c in ln.strip("| ").split("|")]
            if all(re.match(r'^[-:\s]+$', c) for c in cells if c):
                continue
            if any(len(c) > 5000 for c in cells):
                overflow_lines.append(ln)  # 超长行退回当段落
                continue
            rows.append(cells)
        if not rows:
            # 全退回了，当普通段落输出
            for ln in overflow_lines:
                WordExportService._add_body_paragraph(doc, ln.strip('| '))
            return
        max_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols, style='Table Grid')
        # 自适应宽度
        table.autofit = True
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j >= max_cols:
                    break
                cell_text = cell_text.replace('*', '')  # 清除*标记
                cell = table.cell(i, j)
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run()
                # 处理**加粗**和*斜体*
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', cell_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**') and len(part) > 4:
                        r = cell.paragraphs[0].add_run(part[2:-2])
                        r.bold = True
                        _set_run_font(r, FONT_SIZES["body"])
                    elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                        r = cell.paragraphs[0].add_run(part[1:-1])
                        r.italic = True
                        _set_run_font(r, FONT_SIZES["body"])
                    elif part:
                        r = cell.paragraphs[0].add_run(part)
                        _set_run_font(r, FONT_SIZES["body"])
                # 删除第一个空run
                if run.text == '':
                    run._element.getparent().remove(run._element)
                # 表头行加粗
                if i == 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            _set_run_font(r, FONT_SIZES["body"])
        doc.add_paragraph()

    @staticmethod
    def _add_outline_items(
        doc: docx.Document, items: list[WordExportOutlineItem], level: int = 1, first: bool = True
    ) -> None:
        for item in items:
            # 一级章节之间分页（第一个不分）
            if level == 1 and not first:
                doc.add_page_break()
            first = False
            title_text = f"{item.id} {item.title}"

            if level <= 4:
                heading = doc.add_heading(title_text, level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in heading.runs:
                    _set_run_font(run)
            else:
                # 五级及以上：加粗段落
                para = doc.add_paragraph()
                run = para.add_run(title_text)
                run.bold = True
                _set_run_font(run, FONT_SIZES["heading4"])
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(3)

            if not item.children:
                content = item.content or ""
                if content.strip():
                    WordExportService._add_body_content(doc, content)
                continue

            WordExportService._add_outline_items(doc, item.children, level + 1)
